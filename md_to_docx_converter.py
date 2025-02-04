"""
A utility script to convert Markdown files to Microsoft Word (DOCX) format.
This script can process both individual files and entire directories.
"""

# Core Python imports
# Used for command line argument parsing
import argparse
# Used for file and directory operations
import os
# Used for text cleaning and pattern matching in clean_text()
import re

# External dependencies
# Used to convert markdown to HTML
import markdown
# Used to create and manipulate Word documents
from docx import Document
# Used for Word document measurements
from docx.shared import Pt
# Used to parse and traverse HTML content
from bs4 import BeautifulSoup


def clean_text(text):
    """
    Clean and normalize text content.

    Args:
        text (str): Input text to clean

    Returns:
        str: Cleaned text
    """
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    # Remove spaces at the start/end of lines
    text = text.strip()
    # Normalize quotes
    text = text.replace('"', '"').replace('"', '"')
    text = text.replace(''', "'").replace(''', "'")
    # Normalize dashes
    text = text.replace('--', '–')
    # Remove duplicate newlines
    text = re.sub(r'\n\s*\n', '\n', text)
    return text


def process_html_to_docx(soup, doc):
    """
    Process HTML elements and convert them to properly formatted docx elements.

    Args:
        soup (BeautifulSoup): BeautifulSoup object containing HTML
        doc (Document): Word document object
    """
    for element in soup.children:
        if element.name is None:
            # Skip empty text nodes
            if not element.strip():
                continue
            # Add non-empty text as normal paragraph
            text = clean_text(element.string)
            if text:
                p = doc.add_paragraph(text)
                p.style = 'Normal'
            continue

        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            text = clean_text(element.get_text())
            if text:
                p = doc.add_paragraph(text)
                p.style = f'Heading {level}'

        elif element.name == 'p':
            text = clean_text(element.get_text())
            if text:
                p = doc.add_paragraph(text)
                p.style = 'Normal'

        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                text = clean_text(li.get_text())
                if text:
                    p = doc.add_paragraph(text)
                    p.style = 'List Bullet' if element.name == 'ul' else 'List Number'
                    # Ensure proper indentation
                    p.paragraph_format.left_indent = Pt(36)

        elif element.name == 'blockquote':
            text = clean_text(element.get_text())
            if text:
                p = doc.add_paragraph(text)
                p.style = 'Quote'
                # Add left indentation and italics for quotes
                p.paragraph_format.left_indent = Pt(48)
                for run in p.runs:
                    run.italic = True

        elif element.name == 'pre':
            text = clean_text(element.get_text())
            if text:
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                # Add light gray background for code blocks
                p.paragraph_format.left_indent = Pt(36)
                p.paragraph_format.right_indent = Pt(36)

        elif element.name == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.alignment = 1  # 1 = CENTER alignment
            run = p.add_run('─' * 30)
            run.font.size = Pt(12)


def convert_md_to_docx(input_file, output_path=None):
    """
    Convert a markdown file to docx format.

    Args:
        input_file (str): Path to the input markdown file
        output_path (str, optional): Path to output directory for converted files

    Returns:
        bool: True if conversion was successful, False otherwise
    """
    try:
        # Read markdown content
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert markdown to HTML
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')

        # Create a new Word document
        doc = Document()

        # Process the HTML content
        process_html_to_docx(soup, doc)

        # Determine output path
        if output_path is None:
            # Use the same name as input with .docx extension
            output_file = os.path.splitext(input_file)[0] + '.docx'
        else:
            # If output_path is a directory, create output file path inside it
            if os.path.isdir(output_path):
                input_filename = os.path.basename(input_file)
                output_filename = os.path.splitext(input_filename)[0] + '.docx'
                output_file = os.path.join(output_path, output_filename)
            else:
                output_file = output_path

        # Ensure output directory exists
        output_dir = os.path.dirname(output_file)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        # Save the document
        doc.save(output_file)
        print(f"Successfully converted {input_file} to {output_file}")
        return True
    except (IOError, OSError) as e:
        print(f"Error processing file {input_file}: {str(e)}")
        return False


def convert_directory(directory_path, output_path=None):
    """
    Convert all markdown files in a directory to docx format.

    Args:
        directory_path (str): Path to the directory containing markdown files
        output_path (str, optional): Path to output directory for converted files
    """
    if not os.path.isdir(directory_path):
        print(f"Error: {directory_path} is not a directory")
        return

    # Create output directory if specified
    if output_path:
        os.makedirs(output_path, exist_ok=True)

    success_count = 0
    failure_count = 0
    
    # Process all markdown files in the directory
    for filename in os.listdir(directory_path):
        if filename.lower().endswith('.md'):
            input_file = os.path.join(directory_path, filename)
            if convert_md_to_docx(input_file, output_path):
                success_count += 1
            else:
                failure_count += 1

    print("\nConversion complete!")
    print(f"Successfully converted: {success_count} files")
    print(f"Failed conversions: {failure_count} files")


def main():
    """
    Main entry point for the markdown to docx converter.
    Parses command line arguments and initiates the conversion process.
    """
    parser = argparse.ArgumentParser(
        description='Convert markdown files to Microsoft Word format'
    )
    parser.add_argument(
        '--input',
        '-i',
        help='Path to markdown file or directory containing markdown files'
    )
    parser.add_argument(
        '--output',
        '-o',
        help='Output file/directory path (optional)',
        default=None
    )

    args = parser.parse_args()

    if os.path.isfile(args.input):
        # Convert single file
        convert_md_to_docx(args.input, args.output)
    elif os.path.isdir(args.input):
        # Convert all directory
        convert_directory(args.input, args.output)
    else:
        print(f"Error: {args.input} is not a valid file or directory")


if __name__ == "__main__":
    main()
